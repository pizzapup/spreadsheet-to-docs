from flask import (
    Flask,
    request,
    redirect,
    url_for,
    Response,
    render_template_string,
    send_file,
)
import pandas as pd
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
import logging
import zipfile

# Set up logging
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)

# Configure allowed file extensions
ALLOWED_EXTENSIONS = {"xlsx", "csv"}


def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        if "file" not in request.files:
            logging.warning("No file part in the request.")
            return "No file part in the request."
        file = request.files["file"]
        if file.filename == "":
            logging.warning("No file selected.")
            return "No file selected."
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            try:
                # Initialize the file and process it
                df = init_file(file)
                table_html = df.head().to_html(
                    classes="table table-striped", index=False
                )
                column_names = list(df.columns)
                column_feedback = analyze_columns_for_filenames(df)
                return render_template_string(
                    """
                    <!doctype html>
                    <link rel="stylesheet" href="{{ url_for('static', filename='styles.css') }}">
                    <title>File Uploaded</title>
                    <h1>File Uploaded Successfully</h1>
                    <h2>Preview of the File:</h2>
                    {{ table_html|safe }}
                    <form action="/generate_docs" method="post">
                        <input type="hidden" name="data" value="{{ data }}">
                        <label for="filename_template">Filename Template:</label>
                        <input type="text" id="filename_template" name="filename_template" placeholder="e.g., appload-{First Name}-{Last Name}" oninput="showSuggestions(this.value); updateFeedback(this.value);">
                        <div id="suggestions" style="border: 1px solid #ccc; display: none; max-height: 100px; overflow-y: auto;"></div>
                        <br>
                        <div id="column_feedback">
                            <h3>Column Feedback:</h3>
                            <ul id="feedback_list"></ul>
                        </div>
                        <button type="submit">Generate and Download Word Docs</button>
                    </form>
                    <form action="/" method="get">
                        <button type="submit">Upload Another File</button>
                    </form>
                    <script>
                        const columnNames = {{ column_names|tojson }};
                        const columnFeedback = {{ column_feedback|tojson }};
                        
                        function showSuggestions(input) {
                            const suggestionsDiv = document.getElementById('suggestions');
                            suggestionsDiv.innerHTML = '';
                            const lastOpenBraceIndex = input.lastIndexOf('{');
                            if (lastOpenBraceIndex === -1) {
                                suggestionsDiv.style.display = 'none';
                                return;
                            }
                            const prefix = input.substring(lastOpenBraceIndex + 1);
                            const matches = columnNames.filter(col => col.toLowerCase().startsWith(prefix.toLowerCase()));
                            if (matches.length > 0) {
                                matches.forEach(match => {
                                    const suggestion = document.createElement('div');
                                    suggestion.textContent = match;
                                    suggestion.style.cursor = 'pointer';
                                    suggestion.style.padding = '5px';
                                    suggestion.style.borderBottom = '1px solid #ddd';
                                    suggestion.onclick = () => {
                                        const inputField = document.getElementById('filename_template');
                                        inputField.value = inputField.value.substring(0, lastOpenBraceIndex + 1) + match + '}';
                                        suggestionsDiv.style.display = 'none';
                                        updateFeedback(inputField.value);
                                    };
                                    suggestionsDiv.appendChild(suggestion);
                                });
                                suggestionsDiv.style.display = 'block';
                            } else {
                                suggestionsDiv.style.display = 'none';
                            }
                        }

                        function updateFeedback(template) {
                            const feedbackList = document.getElementById('feedback_list');
                            feedbackList.innerHTML = '';
                            const placeholders = [...template.matchAll(/{(.*?)}/g)].map(match => match[1]);
                            const uniquePlaceholders = [...new Set(placeholders)]; // Ensure no duplicates
                            uniquePlaceholders.forEach(placeholder => {
                                if (columnFeedback[placeholder]) {
                                    const feedbackItem = document.createElement('li');
                                    feedbackItem.textContent = columnFeedback[placeholder];
                                    feedbackList.appendChild(feedbackItem);
                                }
                            });
                        }
                    </script>
                    """,
                    table_html=table_html,
                    data=df.to_json(orient="records"),
                    column_names=column_names,
                    column_feedback=column_feedback,
                )
            except ValueError as e:
                logging.error(f"Error processing the file {filename}: {e}")
                return str(e)
            except Exception as e:
                logging.error(f"Unexpected error: {e}")
                return "An unexpected error occurred while processing the file."
        else:
            logging.warning("File type not allowed.")
            return "Invalid file type."
    return """
    <!doctype html>
    <title>Upload New File</title>
    <h1>Upload New File</h1>
    <p>Ensure the spreadsheet includes these columns: "First and Middle Name" and "Last Name".</p>
    <p>Accepted file types: .xlsx, .csv</p>
    <form method=post enctype=multipart/form-data>
      <input type=file name=file>
      <input type=submit value=Upload>
    </form>
    """


def analyze_columns_for_filenames(df):
    """Analyze columns for filename usability and provide feedback."""
    feedback = {}
    for col in df.columns:
        long_values = (
            df[col].apply(lambda x: len(str(x)) > 80 if pd.notnull(x) else False).sum()
        )
        invalid_chars = (
            df[col]
            .apply(
                lambda x: (
                    any(c in str(x) for c in r'\/:*?"<>|') if pd.notnull(x) else False
                )
            )
            .sum()
        )

        if long_values > 0:
            feedback[col] = (
                f"Column '{col}' contains {long_values} value(s) longer than 80 characters. These will be truncated to 60 characters."
            )
            df[col] = df[col].apply(lambda x: str(x)[:60] if pd.notnull(x) else x)

        if invalid_chars > 0:
            feedback[col] = (
                f"Column '{col}' contains {invalid_chars} value(s) with invalid characters. These will be replaced with underscores (_)."
            )
            df[col] = df[col].apply(
                lambda x: (
                    "".join(c if c not in r'\/:*?"<>|' else "_" for c in str(x))
                    if pd.notnull(x)
                    else x
                )
            )

    return feedback


@app.route("/generate_docs", methods=["POST"])
def generate_docs():
    """Generate Word documents from the uploaded data."""
    try:
        data = request.form.get("data")
        filename_template = request.form.get("filename_template", "")
        if not data:
            raise ValueError("No data provided for document generation.")

        df = pd.read_json(data)

        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for i, (_, row) in enumerate(df.iterrows()):
                doc = Document()
                doc.add_heading("Generated Document", level=1)
                for col in df.columns:
                    doc.add_paragraph(f"{col}: {row[col]}")
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Generate filename using the template
                filename = filename_template
                for col in df.columns:
                    placeholder = f"{{{col}}}"
                    filename = filename.replace(placeholder, str(row[col]))
                filename = filename.replace("{index}", str(i))
                if not filename.strip():
                    filename = f"Document_{i}.docx"
                else:
                    filename += ".docx"

                zip_file.writestr(filename, doc_io.read())

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype="application/zip",
            as_attachment=True,
            download_name="documents.zip",
        )
    except Exception as e:
        logging.error(f"Error generating documents: {e}")
        return "An error occurred while generating the documents."


def init_file(file):
    """Initialize the uploaded file and return a DataFrame."""
    try:
        logging.info("Initializing file in memory.")
        if file.filename.endswith(".xlsx"):
            df = pd.read_excel(file)
        elif file.filename.endswith(".csv"):
            df = pd.read_csv(file)
        else:
            raise ValueError("Invalid file type. Only .xlsx and .csv are supported.")

        df.columns = df.columns.str.strip()

        # Check if the required columns exist
        required_columns = {"First and Middle Name", "Last Name"}
        if not required_columns.issubset(df.columns):
            raise ValueError(
                f"Missing required columns. Ensure the file contains: {', '.join(required_columns)}."
            )

        # Check if the DataFrame is empty
        if df.empty:
            raise ValueError("The uploaded file is empty.")

        return df

    except Exception as e:
        logging.error(f"Error initializing file: {e}")
        raise ValueError(
            "Error initializing the file. Please check the file format and content."
        )


if __name__ == "__main__":
    app.run(debug=True)
