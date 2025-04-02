from flask import Blueprint, request, send_file
from docx import Document
from io import BytesIO
import pandas as pd
import zipfile
import logging

generate_blueprint = Blueprint("generate", __name__)


@generate_blueprint.route("/", methods=["POST"])
def generate_docs():
    """Generate Word documents from the uploaded data."""
    try:
        data = request.form.get("data")
        if not data:
            raise ValueError("No data provided for document generation.")

        filename_template = request.form.get("filename_template", "")
        zip_filename = request.form.get("zip_filename", "").strip()
        null_handling = request.form.get("null_handling", "omit")
        null_value = (
            request.form.get("null_value", "N/A") if null_handling == "fill" else None
        )

        df = pd.read_json(data)
        zip_filename = get_zip_filename(zip_filename)
        zip_buffer = generate_zip_file(df, filename_template, null_handling, null_value)
        return send_zip_file(zip_buffer, zip_filename)

    except ValueError as e:
        logging.error(f"Error generating documents: {e}")
        return str(e)


def get_zip_filename(zip_filename):
    if not zip_filename:
        return "Documents.zip"
    elif not zip_filename.lower().endswith(".zip"):
        return zip_filename + ".zip"
    return zip_filename


def handle_null_values(value, null_handling, null_value):
    if pd.isnull(value) or value == "":
        if null_handling == "omit":
            return None
        else:
            return null_value
    return value


def handle_filename_template(filename_template, df, row, i, existing_filenames):
    filename = filename_template
    for col in df.columns:
        placeholder = f"{{{col}}}"
        filename = filename.replace(placeholder, str(row[col]))
    filename = filename.replace("{index}", str(i))
    if not filename.strip():
        filename = f"Document_{i}.docx"
    else:
        filename += ".docx"
    # Ensure unique filenames
    original_filename = filename
    counter = 1
    while filename in existing_filenames:
        filename = f"{original_filename[:-5]}_{counter}.docx"
        counter += 1
    existing_filenames.add(filename)
    return filename


def generate_document(df, row, null_handling, null_value):
    doc = Document()
    doc.add_heading("Generated Document", level=1)
    for col in df.columns:
        value = handle_null_values(row[col], null_handling, null_value)
        if value is not None:
            doc.add_paragraph(f"{col}: {value}")
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def generate_zip_file(df, filename_template, null_handling, null_value):
    zip_buffer = BytesIO()
    existing_filenames = set()  # Track existing filenames
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for i, (_, row) in enumerate(df.iterrows()):
            doc_io = generate_document(df, row, null_handling, null_value)
            filename = handle_filename_template(
                filename_template, df, row, i, existing_filenames
            )
            zip_file.writestr(filename, doc_io.read())
    zip_buffer.seek(0)
    return zip_buffer


def send_zip_file(zip_buffer, zip_filename):
    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name=zip_filename,
    )
