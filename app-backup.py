import os
import logging
import zipfile
from flask import Flask, request, redirect, url_for, send_file, abort, Response
import pandas as pd
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO

# Set up logging
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)

# Configure allowed file extensions
app.config["ALLOWED_EXTENSIONS"] = {"xlsx", "csv"}


def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in app.config["ALLOWED_EXTENSIONS"]
    )


@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        if "file" not in request.files:
            logging.warning("No file part in the request.")
            return "No file part"
        file = request.files["file"]
        if file.filename == "":
            logging.warning("No file selected.")
            return "No selected file"
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            try:
                # Process the file in memory
                return process_file_in_memory(file)
            except Exception as e:
                logging.error(f"Error processing the file {filename}: {e}")
                return "Error processing the file"
        else:
            logging.warning("File type not allowed.")
            return "Invalid file type"
    return """
    <!doctype html>
    <title>Upload new File</title>
    <h1>Upload new File</h1>
    <div>Right now this app requires the spreadsheet to include these columns: "First and Middle Name" and "Last Name"</div>
    <div> Accepted filetypes: xlsx, csv </div>
    <form method=post enctype=multipart/form-data>
      <input type=file name=file>
      <input type=submit value=Upload>
    </form>
    """


def process_file_in_memory(file):
    try:
        logging.info(f"Processing file in memory")

        if file.filename.endswith(".xlsx"):
            df = pd.read_excel(file)
        elif file.filename.endswith(".csv"):
            df = pd.read_csv(file)
        else:
            return "Invalid file type"

        df.columns = df.columns.str.strip()

        # Check if the DataFrame is empty
        if df.empty:
            logging.warning("The file is empty.")
            return "The file is empty."

        # Create a buffer for the ZIP file
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            # Process each row and generate Word documents
            for index, row in df.iterrows():
                doc = Document()

                # Add content to the document
                for col in df.columns:
                    if not pd.isnull(row[col]):
                        question = doc.add_paragraph()
                        question.add_run(f"{col}: ").bold = True
                        doc.add_paragraph(str(row[col]))

                # Save the document to a buffer
                output_buffer = BytesIO()
                doc.save(output_buffer)
                output_buffer.seek(0)

                # Create a filename based on 'Last Name' and 'First and Middle Name'
                filename = f"{row['Last Name']} {row['First and Middle Name']}.docx"

                # Add the Word document to the zip file
                zip_file.writestr(filename, output_buffer.read())

        # Set the pointer to the beginning of the buffer
        zip_buffer.seek(0)

        # Send the ZIP file as a response
        return Response(
            zip_buffer,
            mimetype="application/zip",
            headers={"Content-Disposition": "attachment;filename=processed_files.zip"},
        )

    except Exception as e:
        logging.error(f"Error processing file in memory: {e}")
        return "Error processing the file."


if __name__ == "__main__":
    app.run(debug=True)
